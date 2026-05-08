"""Generate QuantaEX v2 report — QTA Post-Quantum mainnet centric."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def _apply_korean_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = 'Malgun Gothic'
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')


def add_heading(doc, text, level=1, color=(11, 14, 17)):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        _apply_korean_font(run, color=color)
    return h


def add_para(doc, text, bold=False, size=10, color=(40, 40, 40), align=None, italic=False):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    _apply_korean_font(run, size=size, bold=bold, color=color, italic=italic)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    _apply_korean_font(run, size=10)
    return p


def add_table(doc, headers, rows, header_bg='F0B90B'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.autofit = True

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        _apply_korean_font(run, size=10, bold=True, color=(11, 14, 17))
        set_cell_bg(hdr[i], header_bg)

    for r_idx, row_data in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row_data):
            cells[c_idx].text = ''
            p = cells[c_idx].paragraphs[0]
            run = p.add_run(str(val))
            _apply_korean_font(run, size=9.5)
    return table


# ─────────────────────────────────────────────────────────────
doc = Document()

# Base style
style = doc.styles['Normal']
style.font.name = 'Malgun Gothic'
style.font.size = Pt(10)
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = OxmlElement('w:rFonts')
    rPr.append(rFonts)
rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# ── Cover ──────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('QuantaEX 거래소 프로젝트 보고서 v2')
_apply_korean_font(run, size=26, bold=True, color=(240, 185, 11))

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Post-Quantum Native Exchange — QTA 메인넷 중심')
_apply_korean_font(run, size=13, color=(132, 142, 156))

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub2.add_run('세계 최초 양자내성(Post-Quantum) 풀스택 디지털 자산 거래소')
_apply_korean_font(run, size=11, color=(14, 203, 129), bold=True)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('\n작성일: 2026-04-28 (v2)\n도메인: https://quantaex.io\n저장소: github.com/Stevewon/-\n')
_apply_korean_font(run, size=10, color=(100, 100, 100))

doc.add_paragraph()

# ── 0. 변경 사유 ────────────────────────────────────────────
add_heading(doc, '0. v2 개정 사유', level=1)
add_para(doc,
    'v1 보고서는 외부 체인(TRON TRC20) 통합을 다음 단계로 제시했지만, QuantaEX는 '
    '자체 양자내성(Post-Quantum) 메인넷을 보유한 거래소입니다. 따라서 로드맵의 1순위는 '
    '"외부 체인 연동"이 아니라 "QTA 네이티브 메인넷 통합"이며, 외부 체인은 그 위에 '
    '"브릿지(Wrapped QTA)" 형태로 부착되어야 정합성이 맞습니다. 본 v2에서는 이를 전면 반영합니다.',
    size=11)

# ── 1. 한줄 요약 ───────────────────────────────────────────
add_heading(doc, '1. 한줄 요약', level=1)
add_para(doc,
    'QuantaEX는 자체 양자내성 메인넷(QTA Chain) 위에서 동작하는 Cloudflare 엣지 기반 글로벌 '
    'USDT/USDC 거래소입니다. ECDSA 기반 기존 체인이 양자컴퓨팅 시대에 노출되는 자산 탈취 위험을 '
    '원천 차단하며, 거래소 + 메인넷 + 지갑이 모두 Post-Quantum 알고리즘으로 통합 운영되는 '
    '세계 최초의 풀스택 양자내성 거래소를 지향합니다.',
    size=11)

# ── 2. 양자내성 차별화 (NEW — 최상단 배치) ──────────────────
add_heading(doc, '2. 양자내성 차별화 (Post-Quantum Edge)', level=1)
add_para(doc,
    'Shor 알고리즘이 실용화되는 순간 비트코인·이더리움을 포함한 ECDSA/secp256k1 기반 체인의 '
    '지갑 자산은 이론적으로 모두 탈취 가능합니다. QuantaEX QTA 체인은 NIST PQC 표준 알고리즘을 '
    '서명 계층에 채택하여 이 위협으로부터 면역됩니다.',
    size=10)
add_table(doc, ['항목', '기존 거래소 (BTC/ETH)', 'QuantaEX (QTA)'], [
    ['서명 알고리즘', 'ECDSA (secp256k1)', 'NIST PQC 표준 (CRYSTALS-Dilithium 계열)'],
    ['양자컴퓨터 내성', '취약 (Shor 공격)', '내성 (Lattice 기반)'],
    ['지갑 주소 길이', '~34자', '확장 (PQ 공개키 기반)'],
    ['트랜잭션 검증', 'ECDSA 검증', 'PQ 서명 검증 (메인넷 노드 + 거래소 동시 처리)'],
    ['거래소 API 인증', 'HMAC-SHA256', 'PQ 서명 기반 (옵션)'],
])

# ── 3. QTA 메인넷 사양 (NEW) ────────────────────────────────
add_heading(doc, '3. QTA 메인넷 사양', level=1)
add_para(doc, '※ 정식 메인넷 파라미터는 별도 백서/기술문서 기준이며, 본 보고서는 거래소 통합 관점의 요약입니다.', italic=True, size=9, color=(120, 120, 120))
add_table(doc, ['항목', '값/설명'], [
    ['체인 이름', 'QTA Chain (QuantaEX Native Mainnet)'],
    ['네이티브 토큰', 'QTA'],
    ['서명 체계', 'Post-Quantum (NIST PQC 채택)'],
    ['합의 방식', 'PoS 기반 (검증자 모델)'],
    ['블록 타임', '수 초 단위 (실시간 거래소 매칭과 정합)'],
    ['컨펌 정책', '거래소 입금 인정: N 컨펌 (메인넷 안정성 기준)'],
    ['브릿지 가능성', '외부 체인(ETH/BSC) ↔ QTA 양방향 브릿지 (Wrapped QTA, qQTA)'],
])

# ── 4. 기술 스택 ───────────────────────────────────────────
add_heading(doc, '4. 거래소 기술 스택', level=1)
add_table(doc, ['레이어', '기술'], [
    ['프론트엔드', 'React 18 + TypeScript + Vite + TailwindCSS, Zustand'],
    ['차트', 'lightweight-charts (Binance식 캔들·볼륨, 6개 인터벌)'],
    ['백엔드', 'Cloudflare Workers (Hono, SSR 워커)'],
    ['DB', 'Cloudflare D1 (SQLite, 마이그레이션 14건+)'],
    ['스토리지', 'Cloudflare R2 (D1 백업·로그 아카이브)'],
    ['배포', 'Cloudflare Pages + GitHub Actions (Push → 40초 자동 배포)'],
    ['모니터링', 'Sentry + Logflare + 자체 헬스체크'],
    ['부하테스트', 'k6 시나리오'],
    ['체인 통합', 'QTA Mainnet RPC + PQ 서명 검증 라이브러리'],
])

# ── 5. 핵심 기능 ──────────────────────────────────────────
add_heading(doc, '5. 핵심 기능', level=1)

add_heading(doc, '5.1 매칭 엔진 / 주문', level=2)
add_bullet(doc, '시장가 / 지정가 / Stop-Limit (트리거 자동 매칭)')
add_bullet(doc, 'Time-In-Force: IOC / FOK / POST_ONLY')
add_bullet(doc, 'VIP 수수료 티어 + Maker/Taker 분리, append-only fee_ledger')
add_bullet(doc, '시뮬 데이터 100% 제거 — 실제 D1 오더북·체결만 사용')

add_heading(doc, '5.2 자산 / 지갑', level=2)
add_bullet(doc, '글로벌 USDT/USDC 듀얼 베이스 (BTC/ETH/BNB/SOL/XRP/ADA/DOGE/DOT/AVAX/MATIC/QTA × 2 견적)')
add_bullet(doc, 'QTA 네이티브 지갑: PQ 서명 기반 주소 발급, 메인넷 직결')
add_bullet(doc, '사용자별 코인 지갑 자동 생성, 출금 검증 + 2FA 챌린지')

add_heading(doc, '5.3 보안 / 인증', level=2)
add_bullet(doc, 'JWT + localStorage 동기 하이드레이션 (새로고침 로그인 유지)')
add_bullet(doc, '회원가입 1,000 QTA 보너스 (이메일 인증 시 잠금 해제)')
add_bullet(doc, '거래·지갑 이벤트 트랜잭션 이메일')
add_bullet(doc, '회원가입/로그인/출금/비번 변경 시 2FA 챌린지')
add_bullet(doc, '(예정) PQ 서명 기반 API 키 — 양자내성 거래 자동화')

add_heading(doc, '5.4 시장 데이터', level=2)
add_bullet(doc, '실시간 SSE (orderbook, trades, ticker)')
add_bullet(doc, '캔들 차트 6 인터벌 + 5초 라이브 폴링')
add_bullet(doc, '24h 변동/볼륨/고저 자동 집계')

# ── 6. 관리자 콘솔 ────────────────────────────────────────
add_heading(doc, '6. 관리자 콘솔 (PC 최적화)', level=1)
add_para(doc, '240px 고정 사이드바, 11개 탭 그룹화:')
add_table(doc, ['그룹', '탭', '기능'], [
    ['Overview', 'Overview', '사용자/거래/볼륨/수수료/대기 KYC + 14일 트렌드 + Top Markets'],
    ['Operations', 'Users / KYC / Deposits / Withdrawals', 'KYC 승인, 입출금 검토'],
    ['Market', 'Trades / Coins / Broadcast / Fees', '거래 모니터링, 코인 ON/OFF, 공지, 수수료 원장'],
    ['System', 'Audit / System', '감사 로그, 마이그레이션·백업·헬스체크, (예정) QTA 노드 상태'],
])

# ── 7. 옵저버빌리티 ───────────────────────────────────────
add_heading(doc, '7. 옵저버빌리티 / DevOps', level=1)
add_bullet(doc, 'Sentry — 에러 추적')
add_bullet(doc, 'Logflare — 구조화 로그 수집')
add_bullet(doc, 'D1 → R2 일일 자동 백업 (JSONL 압축, 30일 보존)')
add_bullet(doc, '헬스체크 — /api/health, /api/admin/system-status')
add_bullet(doc, 'k6 부하 테스트 — 가입·로그인·주문·출금 풀 저니')
add_bullet(doc, 'GitHub Actions CI/CD — 40초 이내 자동 배포')

# ── 8. 성능 지표 ──────────────────────────────────────────
add_heading(doc, '8. 성능 지표 (Sprint 4 직전 측정)', level=1)
add_table(doc, ['항목', '수치'], [
    ['/api/market/markets', '0.15s'],
    ['/api/market/tickers (콜드)', '0.30s (N+1 → 3쿼리, 3배 개선)'],
    ['/api/market/tickers (웜)', '0.28s'],
    ['마켓 페이지 첫 화면', '<1초 (이전 8~15초 → 10배 개선)'],
    ['캔들 차트 초기 로드', '<0.5초 + 5초 라이브 폴링'],
    ['Cloudflare Pages 빌드', '평균 38~44초'],
])

# ── 9. 마일스톤 ───────────────────────────────────────────
add_heading(doc, '9. 마일스톤 (Sprint 3 → Sprint 4)', level=1)
add_table(doc, ['Sprint', '핵심 성과'], [
    ['Sprint 3', '시뮬 제거 / 어드민 감사 로그 / Stop-Limit / TIF / VIP 수수료 / 거래 이메일 / 가입 보너스 잠금'],
    ['Sprint 3+', '옵저버빌리티 / D1→R2 백업 / k6 / Audit·Fees 탭 / System 헬스 / PC 어드민 사이드바'],
    ['Sprint 4 Phase A ✅', 'KRW 제거 + USDC 추가 + 글로벌화 + Markets 10배 가속 + 차트 setInterval 셰도잉 버그 수정'],
])

# ── 10. v2 로드맵 (QTA 메인넷 중심) ────────────────────────
add_heading(doc, '10. v2 로드맵 — QTA 메인넷 중심', level=1)
add_para(doc, '거래소 완성 순서 (외부 체인보다 자체 메인넷 우선):', bold=True)
add_table(doc, ['Phase', '주제', '핵심 작업'], [
    ['Phase B ⏳', 'QTA 네이티브 입출금 통합',
     'QTA Chain RPC 연동, PQ 서명 검증, 사용자별 QTA 주소 발급, 입금 모니터링 cron, 컨펌 로직, 핫/콜드 월렛 분리, 출금 큐 + 멀티 서명'],
    ['Phase C', 'QTA 운영 어드민 + Risk',
     '검증자 노드 상태, 블록 높이, 메모리풀, 핫월렛 잔고, Withdrawal Queue, IP 차단/2FA 강제 해제 감사, 마켓 서킷 브레이커'],
    ['Phase D', '브릿지: QTA ↔ ETH (Wrapped qQTA)',
     '락-앤-민트/번-앤-언락 컨트랙트, 브릿지 검증자, qQTA(ERC-20) 발행, 외부 유동성 확보'],
    ['Phase E', '외부 체인 입금 브릿지 확장',
     'ETH·USDT·USDC를 QTA 메인넷으로 가져오기 (역방향 브릿지), BSC 추가'],
    ['Phase F', '선물(Perpetual) / 마진',
     'USDT-M Perp, 펀딩비, 청산 엔진, 격리/교차 마진'],
    ['Phase G', 'PQ API Trading + WebSocket',
     '양자내성 서명 기반 API 키, 공식 WebSocket 키 발급, REST/WS Rate Limit'],
])

# ── 11. 차별화 포인트 ─────────────────────────────────────
add_heading(doc, '11. 차별화 포인트 (요약)', level=1)
diff_points = [
    '세계 최초 풀스택 양자내성 거래소 — 거래소 + 메인넷 + 지갑 통합 PQ',
    '엣지 네이티브 풀스택 — 별도 인프라·k8s 없이 글로벌 200+ POP에서 ms 응답',
    'append-only fee_ledger — 수수료 변조 불가, 회계 감사 친화',
    '시뮬 데이터 0% — 모든 가격·체결이 실제 D1 데이터',
    'PC 최적화 어드민 — 240px 사이드바, 1600px max-width, 모바일 드로어 동시 지원',
    '자동 백업 + 헬스체크 + 감사 로그 — 운영 책임 추적성 확보',
    '글로벌 USD 단일 표기 — 환산 환율 의존 0, 환위험 0',
    'VIP 티어 자동 적용 — 30일 거래량 기반 자동 강등/승급',
]
for p in diff_points:
    add_bullet(doc, p)

# ── 12. 결론 ──────────────────────────────────────────────
add_heading(doc, '12. 결론', level=1)
add_para(doc,
    'QuantaEX는 "자체 양자내성 메인넷(QTA) + 풀스펙 거래소 + 엣지 네이티브 성능 + 운영 친화 어드민"의 '
    '4축을 모두 갖춘 차세대 거래소입니다. v1 로드맵에 있던 외부 체인 우선 통합 노선을 폐기하고, '
    '자체 메인넷 통합(Phase B) → 운영 어드민(Phase C) → 외부 체인 브릿지(Phase D~E) → 파생상품(Phase F) → '
    'PQ API(Phase G) 순으로 거래소를 완성합니다. 이는 양자컴퓨팅 시대에 자산을 안전하게 보관할 수 있는 '
    '거의 유일한 글로벌 거래소 포지션을 확보하는 길입니다.',
    size=11)

# Footer
doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = foot.add_run('— QuantaEX Project Report v2 · Post-Quantum Native Exchange · 2026-04-28 —')
_apply_korean_font(run, size=9, color=(150, 150, 150), italic=True)

out_path = '/home/user/webapp/reports/QuantaEX_Report_v2_2026-04-28.docx'
doc.save(out_path)
print(f'✅ Saved: {out_path}')
