"""Generate QuantaEX project report as a Word (.docx) file."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color=(11, 14, 17)):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(*color)
        run.font.name = 'Malgun Gothic'
        # Set East-Asian font
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    return h


def add_para(doc, text, bold=False, size=10, color=(40, 40, 40), align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*color)
    run.font.name = 'Malgun Gothic'
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Malgun Gothic'
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    return p


def add_table(doc, headers, rows, header_bg='F0B90B'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.autofit = True

    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(11, 14, 17)
        run.font.size = Pt(10)
        run.font.name = 'Malgun Gothic'
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
        set_cell_bg(hdr[i], header_bg)

    # Body rows
    for r_idx, row_data in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row_data):
            cells[c_idx].text = ''
            p = cells[c_idx].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)
            run.font.name = 'Malgun Gothic'
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    return table


# ─────────────────────────────────────────────────────────────
doc = Document()

# Set base document font
style = doc.styles['Normal']
style.font.name = 'Malgun Gothic'
style.font.size = Pt(10)
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = OxmlElement('w:rFonts')
    rPr.append(rFonts)
rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# ── Cover ──────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('QuantaEX 거래소 프로젝트 보고서')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(240, 185, 11)  # Binance yellow
run.font.name = 'Malgun Gothic'
rPr = run._element.get_or_add_rPr()
rFonts = rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = OxmlElement('w:rFonts')
    rPr.append(rFonts)
rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Cloudflare Edge 기반 글로벌 디지털 자산 거래소')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(132, 142, 156)
run.font.name = 'Malgun Gothic'

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('\n작성일: 2026-04-28\n도메인: https://quantaex.io\n저장소: github.com/Stevewon/-\n')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(100, 100, 100)
run.font.name = 'Malgun Gothic'

doc.add_paragraph()

# ── 1. 요약 ────────────────────────────────────────────────
add_heading(doc, '1. 한줄 요약', level=1)
add_para(doc,
    'QuantaEX는 Cloudflare 엣지(Workers + D1 + R2) 위에서 풀스택으로 동작하는 글로벌 USDT/USDC '
    '기반 디지털 자산 거래소입니다. KRW를 제거하고 USD 표기로 통일한 글로벌 노선이며, 자체 매칭 엔진· '
    'VIP 수수료 구조·관리자 콘솔·옵저버빌리티·자동 백업까지 갖춘 풀스펙 거래소 플랫폼입니다.',
    size=11)

# ── 2. 기술 스택 ───────────────────────────────────────────
add_heading(doc, '2. 기술 스택', level=1)
add_table(doc, ['레이어', '기술'], [
    ['프론트엔드', 'React 18 + TypeScript + Vite + TailwindCSS, Zustand 상태관리'],
    ['차트', 'lightweight-charts (Binance식 캔들·볼륨 차트, 6개 인터벌)'],
    ['백엔드', 'Cloudflare Workers (Hono 라우팅, SSR 워커)'],
    ['DB', 'Cloudflare D1 (SQLite, 마이그레이션 14건+)'],
    ['스토리지', 'Cloudflare R2 (D1 백업·로그 아카이브)'],
    ['배포', 'Cloudflare Pages + GitHub Actions (Push → 40초 자동 배포)'],
    ['모니터링', 'Sentry + Logflare + 자체 헬스체크 시스템'],
    ['부하테스트', 'k6 시나리오'],
])
add_para(doc, '핵심: 콜드 스타트 0ms, 글로벌 200+ 엣지에서 즉시 응답 — 전통적 클라우드 대비 응답 지연 1/10 수준.',
         bold=True, size=10, color=(14, 203, 129))

# ── 3. 핵심 기능 ──────────────────────────────────────────
add_heading(doc, '3. 핵심 기능', level=1)

add_heading(doc, '3.1 매칭 엔진 / 주문 시스템', level=2)
add_bullet(doc, '시장가 / 지정가 / Stop-Limit 주문 (트리거 가격 도달 시 자동 매칭)')
add_bullet(doc, 'Time-In-Force: IOC / FOK / POST_ONLY 옵션 지원')
add_bullet(doc, 'VIP 수수료 티어 + Maker/Taker 분리, append-only fee_ledger')
add_bullet(doc, '시뮬 데이터 100% 제거 → 실제 D1 오더북·체결 내역만 사용')

add_heading(doc, '3.2 자산 / 지갑', level=2)
add_bullet(doc, '글로벌 USDT/USDC 듀얼 베이스 페어 (BTC, ETH, BNB, SOL, XRP, ADA, DOGE, DOT, AVAX, MATIC, QTA — 11개 코인 × 2 견적)')
add_bullet(doc, '사용자별 코인 지갑 자동 생성, 출금 검증 + 2FA 챌린지')
add_bullet(doc, 'KRW 완전 제거 (Sprint 4 Phase A) — 글로벌 거래소 노선 확정')

add_heading(doc, '3.3 보안 / 인증', level=2)
add_bullet(doc, 'JWT 토큰 + localStorage 동기 하이드레이션 (새로고침 시 로그인 유지)')
add_bullet(doc, '회원가입 1,000 QTA 보너스 (이메일 인증 시 잠금 해제)')
add_bullet(doc, '거래·지갑 이벤트 트랜잭션 이메일 알림')
add_bullet(doc, '회원가입 / 로그인 / 출금 / 비밀번호 변경 시 2FA 챌린지')

add_heading(doc, '3.4 시장 데이터', level=2)
add_bullet(doc, '실시간 SSE 스트림 (orderbook, trades, ticker)')
add_bullet(doc, '캔들 차트 6 인터벌 (1m/5m/15m/1h/4h/1d) + 5초 폴링 라이브 업데이트')
add_bullet(doc, '24h 변동률·볼륨·고/저가 자동 집계')

add_heading(doc, '3.5 가격 알림', level=2)
add_bullet(doc, '사용자 정의 가격 트리거 → 도달 시 푸시/이메일 알림')
add_bullet(doc, '관리자 강제 알림 체크 트리거 가능')

# ── 4. 관리자 콘솔 ────────────────────────────────────────
add_heading(doc, '4. 관리자 콘솔 (PC 최적화)', level=1)
add_para(doc, 'PC 최적화 사이드바 레이아웃 (240px 고정), 11개 탭 그룹화:')
add_table(doc, ['그룹', '탭', '기능'], [
    ['Overview', 'Overview', '사용자/거래/볼륨/수수료/대기 KYC 등 통계 + 14일 트렌드 차트 + Top Markets'],
    ['Operations', 'Users / KYC / Deposits / Withdrawals', 'KYC 승인, 입출금 검토'],
    ['Market', 'Trades / Coins / Broadcast / Fees', '거래 모니터링, 코인 ON/OFF, 공지 푸시, 수수료 원장'],
    ['System', 'Audit / System', '모든 어드민 액션 감사 로그, DB 마이그레이션·백업·헬스체크'],
])
add_para(doc, '특히 System 탭: 6개 시스템 마커(테이블·인덱스·컬럼) 자동 검증, R2 마지막 백업 시각, '
              'Sentry/Logflare 연결 상태를 한눈에 확인 가능.', size=10)

# ── 5. 옵저버빌리티 ───────────────────────────────────────
add_heading(doc, '5. 옵저버빌리티 / DevOps', level=1)
add_bullet(doc, 'Sentry — 에러 추적')
add_bullet(doc, 'Logflare — 구조화 로그 수집')
add_bullet(doc, 'D1 → R2 일일 자동 백업 — JSONL 압축, 30일 보존')
add_bullet(doc, '헬스체크 엔드포인트 — /api/health, /api/admin/system-status')
add_bullet(doc, 'k6 부하 테스트 시나리오 — 회원가입·로그인·주문·출금 풀 저니')
add_bullet(doc, 'GitHub Actions CI/CD — 빌드·마이그레이션·Pages 배포 40초 이내')

# ── 6. 성능 지표 ──────────────────────────────────────────
add_heading(doc, '6. 성능 지표 (Sprint 4 직전 측정)', level=1)
add_table(doc, ['항목', '수치'], [
    ['/api/market/markets', '0.15s'],
    ['/api/market/tickers (콜드)', '0.30s (N+1 → 3쿼리, 3배 개선)'],
    ['/api/market/tickers (웜)', '0.28s'],
    ['마켓 페이지 첫 화면', '<1초 (이전 8~15초 → 10배 개선)'],
    ['캔들 차트 초기 로드', '<0.5초 + 5초 라이브 폴링'],
    ['Cloudflare Pages 빌드', '평균 38~44초'],
])

# ── 7. 마일스톤 ───────────────────────────────────────────
add_heading(doc, '7. 최근 주요 마일스톤 (Sprint 3 → Sprint 4)', level=1)
add_table(doc, ['Sprint', '핵심 성과'], [
    ['Sprint 3', '시뮬 데이터 제거 / 어드민 감사 로그 / Stop-Limit / TIF / VIP 수수료 / 거래 이메일 / 가입 보너스 잠금'],
    ['Sprint 3+', '옵저버빌리티(Sentry/Logflare) / D1→R2 백업 / k6 부하 테스트 / Audit·Fees 탭 / System 헬스 / PC 어드민 사이드바'],
    ['Sprint 4 Phase A ✅', 'KRW 완전 제거 + USDC 추가 + 글로벌화 + Markets 페이지 10배 가속 + 차트 setInterval 셰도잉 버그 수정'],
    ['Sprint 4 Phase B', 'TRON TRC20 USDT/USDC 입출금 PoC (Shasta 테스트넷) — 진행 예정'],
    ['Sprint 4 Phase C', '어드민 추가: Wallets·Withdrawal Queue·Risk 탭 — 진행 예정'],
    ['Sprint 4 Phase D', 'ERC-20 / BEP-20 확장 — 진행 예정'],
])

# ── 8. 차별화 ─────────────────────────────────────────────
add_heading(doc, '8. 차별화 포인트', level=1)
diff_points = [
    ('엣지 네이티브 풀스택', '별도 인프라·k8s 없이 글로벌 200+ POP에서 ms 응답'),
    ('append-only fee_ledger', '수수료 변조 불가능, 회계 감사 친화'),
    ('시뮬 데이터 0%', '모든 가격·체결이 실제 D1 데이터'),
    ('PC 최적화 어드민', '240px 사이드바, 1600px max-width, 모바일 드로어 동시 지원'),
    ('자동 백업 + 헬스체크 + 감사 로그', '운영 책임 추적성 확보'),
    ('글로벌 USD 단일 표기', '환산 환율(×1350) 의존 0, 환위험 0'),
    ('VIP 티어 자동 적용', '30일 거래량 기반 자동 강등/승급'),
]
for i, (k, v) in enumerate(diff_points, 1):
    add_bullet(doc, f'{k}: {v}')

# ── 9. 로드맵 ─────────────────────────────────────────────
add_heading(doc, '9. 향후 로드맵 요약', level=1)
add_bullet(doc, 'Phase B (다음): TRON TRC20 USDT/USDC 실제 입출금 (HD 지갑, 5분 모니터링 cron, 19 confirmations)')
add_bullet(doc, 'Phase C: Wallets / Withdrawal Queue / Risk 탭 어드민 확장')
add_bullet(doc, 'Phase D: Ethereum ERC-20 + BSC BEP-20 다중 체인')
add_bullet(doc, 'Phase E (계획): 선물(Perpetual) / 마진 거래')
add_bullet(doc, 'Phase F (계획): API Trading + WebSocket 공식 키 발급')

# ── 10. 결론 ──────────────────────────────────────────────
add_heading(doc, '10. 결론', level=1)
add_para(doc,
    'QuantaEX는 "전통 거래소 풀스펙 + 엣지 네이티브 성능 + 운영 친화 어드민" 세 축을 모두 만족하는 거래소입니다. '
    'Sprint 3까지 거래소 코어 기능 완성, Sprint 4 Phase A에서 글로벌화 완료, Phase B부터는 진짜 블록체인 입출금이 '
    '붙으면서 테스트넷 거래소 → 실거래 글로벌 거래소로 전환되는 단계입니다.',
    size=11)

# Footer
doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = foot.add_run('— QuantaEX Project Report · 2026-04-28 —')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(150, 150, 150)
run.italic = True
run.font.name = 'Malgun Gothic'

out_path = '/home/user/webapp/reports/QuantaEX_Report_2026-04-28.docx'
doc.save(out_path)
print(f'✅ Saved: {out_path}')
